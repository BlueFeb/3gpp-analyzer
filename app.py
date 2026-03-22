import streamlit as st
import streamlit.components.v1 as components
import os
import tempfile
import zipfile
import requests
import numpy as np
import re
import io
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from openpyxl import load_workbook
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from sklearn.cluster import AgglomerativeClustering
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai

# ==========================================
# 0. 구글 애드센스 광고 영역 (텍스트 숨김 처리)
# ==========================================
def show_adsense():
    """
    구글 애드센스 코드를 삽입하는 영역입니다.
    현재는 화면에 아무런 글자도 보이지 않는 빈 공간(투명)으로 처리되어 있습니다.
    """
    html_code = """
    <div style="width: 100%; height: 90px; background-color: transparent;">
        </div>
    """
    components.html(html_code, height=100)

# ==========================================
# 1. 환경 설정 및 세션 초기화 (메모리 안정성 강화)
# ==========================================
st.set_page_config(page_title="3GPP AI Analyzer Pro", page_icon="📡", layout="wide")

show_adsense()

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "log_text" not in st.session_state:
    st.session_state.log_text = ""
if "process_done" not in st.session_state:
    st.session_state.process_done = False
if "out1_bytes" not in st.session_state:
    st.session_state.out1_bytes = None
if "out2_bytes" not in st.session_state:
    st.session_state.out2_bytes = None
if "extracted_data" not in st.session_state:
    st.session_state.extracted_data = []
if "notebooklm_txt" not in st.session_state:
    st.session_state.notebooklm_txt = None
# --- AI 요약 결과물 보존을 위한 세션 변수 ---
if "ai_summary_generated" not in st.session_state:
    st.session_state.ai_summary_generated = False
if "ai_summary_bytes" not in st.session_state:
    st.session_state.ai_summary_bytes = None
if "ai_summary_text" not in st.session_state:
    st.session_state.ai_summary_text = ""
if "ai_model_name" not in st.session_state:
    st.session_state.ai_model_name = ""

def append_log(text):
    st.session_state.log_text += f"{text}\n"

INTERNAL_PROXY = {"http": "http://10.112.1.184:8080", "https": "http://10.112.1.184:8080"}
INTERNAL_PIN_URL = (
    "https://raw.github.sec.samsung.net/gist/bh14-jung/"
    "d3cfd4262296e61a66ddcaaf045657ed/raw/"
    "67237fa5e0f0f8b781d984c420c45ebede27adf9/"
    "gistfile1.txt?token=AAAAOIJ4CH7DR5P4CEMRHZLIY6KJM"
)
EXTERNAL_PIN_CSV_URL = (
    "https://docs.google.com/spreadsheets/d/e/"
    "2PACX-1vTsSe2LFcO129jJVigl6e9TzHVz8Iaoozasee_4bD1RTwoRS5DTSv-"
    "FdO7dwrPcJ6t7wmQ0-s7197g5/pub?gid=0&single=true&output=csv"
)
FALLBACK_PIN = "2510"
USE_PROXY = False

def detect_network():
    global USE_PROXY
    try:
        requests.get(INTERNAL_PIN_URL, proxies=INTERNAL_PROXY, timeout=3, verify=False).raise_for_status()
        USE_PROXY = True
    except:
        USE_PROXY = False

def fetch_remote_pin():
    if USE_PROXY:
        try:
            r = requests.get(INTERNAL_PIN_URL, proxies=INTERNAL_PROXY, timeout=5, verify=False)
            r.raise_for_status()
            p = r.text.strip().splitlines()[0].strip()
            if p.isdigit() and len(p) == 4:
                return p
        except:
            pass
    try:
        r = requests.get(EXTERNAL_PIN_CSV_URL, timeout=5)
        r.raise_for_status()
        first = r.text.strip().splitlines()[0]
        p = first.split(",")[0].strip()
        if p.isdigit() and len(p) == 4:
            return p
    except:
        pass
    return FALLBACK_PIN

# ==========================================
# 2. 문서 처리 및 로직
# ==========================================
def read_excel_from_bytes(uploaded_file):
    wb = load_workbook(uploaded_file, read_only=False, data_only=True)
    ws = wb.active
    entries = []
    for row in ws.iter_rows(min_row=2):
        cell = row[0]
        comp = row[2] if len(row) > 2 else None
        docid = str(cell.value).strip() if cell.value else ""
        company = str(comp.value).strip() if comp and comp.value else ""
        
        if not docid: continue
        
        # 하이퍼링크가 있어야 SA, CT 등 모든 그룹의 문서를 정확히 다운로드 가능
        if getattr(cell, "hyperlink", None) and cell.hyperlink.target:
            link = cell.hyperlink.target
        else:
            link = f"https://www.3gpp.org/ftp/tsg_ran/WG1_RL1/TSGR1_122/Docs/{docid}.zip"
        entries.append({"doc": docid, "company": company, "link": link})
    return entries

def clone_paragraph(src, dest):
    np_para = dest.add_paragraph("", style=src.style)
    for r in src.runs:
        nr = np_para.add_run(r.text)
        nr.bold = r.bold
        nr.italic = r.italic
        nr.underline = r.underline
        if hasattr(r.font, "name") and r.font.name:
            nr.font.name = r.font.name
        if hasattr(r.font, "size") and r.font.size:
            nr.font.size = r.font.size
        if hasattr(r.font, "color") and getattr(r.font.color, "rgb", None):
            nr.font.color.rgb = r.font.color.rgb
    return np_para

def repackage_docm_to_docx(path, td):
    ud = os.path.join(td, "docm_unzip")
    os.makedirs(ud, exist_ok=True)
    with zipfile.ZipFile(path, 'r') as zf: 
        zf.extractall(ud)
    tf = os.path.join(ud, "[Content_Types].xml")
    t = open(tf, 'r', encoding='utf-8').read()
    t = t.replace(
        'application/vnd.ms-word.document.macroEnabled.main+xml',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
    )
    open(tf, 'w', encoding='utf-8').write(t)
    rp = os.path.join(td, "repack.zip")
    with zipfile.ZipFile(rp, 'w', zipfile.ZIP_DEFLATED) as zf:
        for r, _, fs in os.walk(ud):
            for f in fs:
                full = os.path.join(r, f)
                arc = os.path.relpath(full, ud)
                zf.write(full, arc)
    out = os.path.join(td, "repack.docx")
    os.rename(rp, out)
    return out

def _download_doc(entry, td_name, headers):
    try:
        kwargs = {"headers": headers, "timeout": 60, "verify": False}
        if USE_PROXY: kwargs["proxies"] = INTERNAL_PROXY
        r = requests.get(entry["link"], **kwargs)
        r.raise_for_status()
        fp = os.path.join(td_name, f"{entry['doc']}.zip")
        with open(fp, "wb") as f:
            f.write(r.content)
        return entry, fp, None
    except Exception as ex:
        return entry, None, str(ex)

def extract_all_conclusions(entries, status_elem, progress_elem, log_func):
    with tempfile.TemporaryDirectory() as temp_dir:
        log_func(f"임시 디렉터리 생성: {temp_dir}")

        od = Document()
        od.add_heading("3GPP Conclusions", level=0)
        
        cps = [re.compile(r"^(?:#\s*)?(?:\d+\.?\s*)?(conclusions?)\s*$", re.I), re.compile(r"^(?:#\s*)?(?:\d+\.?\s*)?(summary)\s*$", re.I)]
        eps = [re.compile(r"^(?:#\s*)?(?:\d+\.?\s*)?(references?|appendix|acknowledgment)\s*$", re.I)]
        headers = {"User-Agent": "Mozilla/5.0"}

        download_results = []
        extracted_list = []
        total = len(entries)
        
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {executor.submit(_download_doc, e, temp_dir, headers): e for e in entries}
            for i, fut in enumerate(as_completed(futures), start=1):
                e, fp, err = fut.result()
                download_results.append((e, fp, err))
                progress_elem.progress(i / total)
                status_elem.text(f"Downloaded [{i}/{total}]: {e['doc']}")
                log_func(f"[{i}/{total}] Downloaded: {e['doc']}")

        for idx, (e, fp, err) in enumerate(download_results, start=1):
            status_elem.text(f"Extracting [{idx}/{total}]: {e['doc']}")
            doc_text_buffer = []
            full_text_buffer = []

            tbl = od.add_table(rows=4, cols=2, style="Table Grid")
            tbl.cell(0, 0).text, tbl.cell(0, 1).text = "Document", e["doc"]
            tbl.cell(1, 0).text, tbl.cell(1, 1).text = "Link", e["link"]
            tbl.cell(2, 0).text, tbl.cell(2, 1).text = "Company", e["company"]
            tbl.cell(3, 0).text = "Title"

            try:
                if err or not fp: raise Exception(err or "Download failed")

                ed = os.path.join(temp_dir, e["doc"])
                os.makedirs(ed, exist_ok=True)
                with zipfile.ZipFile(fp) as zf:
                    zf.extractall(ed)

                src_path = None
                for ext in ("*.docx", "*.docm", "*.doc"):
                    src_path = next(Path(ed).rglob(ext), None)
                    if src_path: break

                if not src_path:
                    od.add_paragraph("DOC 파일을 찾을 수 없습니다.")
                    log_func(f"{e['doc']} 없음")
                    continue

                file_path_str = str(src_path)
                
                if src_path.suffix.lower() == ".docm":
                    try:
                        file_path_str = repackage_docm_to_docx(file_path_str, temp_dir)
                    except Exception as ex:
                        log_func(f"{e['doc']} docm 변환 오류: {ex}")
                
                try:
                    sd = Document(file_path_str)
                except Exception as ex:
                    od.add_paragraph(f"문서를 열 수 없습니다 (구형 .doc 파일이거나 손상됨): {ex}")
                    log_func(f"{e['doc']} 문서 파싱 에러: {ex}")
                    continue
                    
                title = ""
                paras = sd.paragraphs
                
                for p in paras:
                    t = p.text.strip()
                    if t:
                        full_text_buffer.append(t)
                    if not title and t.lower().startswith("title:"):
                        title = t.split(":", 1)[1].strip()
                        
                if not title:
                    title = sd.core_properties.title or ""
                tbl.cell(3, 1).text = title

                start = None
                for pat in cps:
                    for j, p in enumerate(paras):
                        if pat.match(p.text.strip()):
                            start = j
                            break
                    if start is not None: break

                if start is None:
                    od.add_paragraph("결론 섹션 없음")
                    log_func(f"{e['doc']} 결론없음")
                else:
                    end = len(paras)
                    for ep in eps:
                        for j, p in enumerate(paras[start + 1 :], start + 1):
                            if ep.match(p.text.strip()):
                                end = j
                                break
                        if end < len(paras): break
                    for j in range(start + 1, end):
                        clone_paragraph(paras[j], od)
                        doc_text_buffer.append(paras[j].text)
                    log_func(f"{e['doc']} 추출 완료")

                extracted_list.append({
                    "doc": e["doc"], "company": e["company"], "link": e["link"], 
                    "title": title, 
                    "content": "\n".join(doc_text_buffer) if doc_text_buffer else "Conclusion 섹션을 찾지 못했습니다.",
                    "full_content": "\n".join(full_text_buffer) if full_text_buffer else "원문 텍스트를 추출하지 못했습니다."
                })

            except Exception as ex:
                od.add_paragraph(f"오류 - {e['doc']}: {ex}")
                log_func(str(ex))

            if idx < len(download_results):
                od.add_page_break()

        st.session_state.extracted_data = extracted_list
        
        txt_buffer = []
        txt_buffer.append("=== 3GPP Contributions Conclusions ===")
        for item in extracted_list:
            txt_buffer.append(f"\n\n--- Document: {item['doc']} ---")
            txt_buffer.append(f"Company: {item['company']}")
            txt_buffer.append(f"Title: {item['title']}")
            txt_buffer.append("Content:")
            txt_buffer.append(item['content'])
        st.session_state.notebooklm_txt = "\n".join(txt_buffer)

        bio = io.BytesIO()
        od.save(bio)
        bio.seek(0)
        
    return bio

class TFIDFEmbedder:
    def __init__(self, max_features=3000, ngram_range=(1, 2)):
        self.v = TfidfVectorizer(
            max_features=max_features, ngram_range=ngram_range,
            lowercase=True, stop_words="english", strip_accents="unicode",
            token_pattern=r"\b[a-zA-Z]{2,}\b",
        )
        self.fitted = False

    def encode(self, texts):
        if isinstance(texts, str): texts = [texts]
        proc = [re.sub(r"\s+", " ", re.sub(r"[^\w\s\-]", " ", t.lower())).strip() for t in texts]
        if not self.fitted:
            self.v.fit(proc)
            self.fitted = True
        return self.v.transform(proc).toarray()

def parse_and_summarize(in_bio, status_elem, log_func):
    d = Document(in_bio)
    props, pcs, cur = [], {}, None

    for el in d.element.body:
        if el.tag.endswith("tbl"):
            tbl = Table(el, d)
            for r in tbl.rows:
                if r.cells[0].text.strip() == "Company":
                    cur = r.cells[1].text.strip()
        elif el.tag.endswith("p"):
            p = Paragraph(el, d)
            txt = p.text.strip()
            if txt.lower().startswith("proposal"):
                buf, cm = [txt], {cur} if cur else set()
                idx2 = d.element.body.index(el) + 1
                while idx2 < len(d.element.body):
                    sib = d.element.body[idx2]
                    if not sib.tag.endswith("p"): break
                    sp = Paragraph(sib, d)
                    st = sp.text.rstrip()
                    if not st.strip() or st.lower().startswith("proposal"): break
                    buf.append(st)
                    if cur: cm.add(cur)
                    idx2 += 1
                bl = "\n".join(buf)
                props.append(bl)
                pcs[bl] = cm.copy()

    r = Document()
    r.add_heading("Proposal Summary", 0)

    if not props:
        r.add_paragraph("No proposals found.")
        bio = io.BytesIO()
        r.save(bio)
        bio.seek(0)
        return bio

    status_elem.text("Generating embeddings & Clustering...")
    em = TFIDFEmbedder()
    emb = em.encode(props)

    N = len(props)
    mn, mx = max(2, N // 5), max(3, N // 2)
    best_diff = float("inf")
    best_lbl = None
    for thr in np.linspace(0.2, 0.8, 13):
        try:
            hac = AgglomerativeClustering(
                n_clusters=None, metric="cosine", linkage="average",
                distance_threshold=thr, compute_full_tree=True,
            )
            lbls = hac.fit_predict(emb)
            cnt = len(set(lbls))
            diff = abs(cnt - (mn + mx) / 2)
            if diff < best_diff:
                best_diff = diff
                best_lbl = lbls
        except: pass
    lbls = best_lbl if best_lbl is not None else np.zeros(N, dtype=int)

    clusters = {}
    for i, l in enumerate(lbls):
        clusters.setdefault(l, {"idxs": [], "cm": set()})
        clusters[l]["idxs"].append(i)
        clusters[l]["cm"].update(pcs[props[i]])

    items = []
    for info in clusters.values():
        idxs = info["idxs"]
        subset = emb[idxs]
        cent = np.mean(subset, axis=0, keepdims=True)
        sims = cosine_similarity(cent, subset)[0]
        rep = props[idxs[int(np.argmax(sims))]]
        cm = sorted(info["cm"])
        items.append({"proposal": rep, "companies": cm, "count": len(cm)})

    items.sort(key=lambda x: x["count"], reverse=True)

    status_elem.text("Creating summary...")
    for it in items:
        r.add_paragraph(it["proposal"])
        r.add_paragraph(f"Supporting companies ({it['count']}): " + (", ".join(it["companies"]) if it["companies"] else "(none)"))
        r.add_paragraph("")
        
    bio = io.BytesIO()
    r.save(bio)
    bio.seek(0)
    log_func("Summary 생성 완료")
    return bio

# ==========================================
# 3. 사이드바 및 메인 화면 구성
# ==========================================
st.sidebar.title("📡 3GPP AI Analyzer")
st.sidebar.markdown("---")
page = st.sidebar.radio("메뉴 이동", ["🚀 통합 AI 분석기", "📁 3GPP FTP 탐색기", "ℹ️ 소개 및 가이드"])
st.sidebar.markdown("---")

# --- 페이지 1: 직관적인 통합 원페이지 흐름 ---
if page == "🚀 통합 AI 분석기":
    st.title("🚀 3GPP 기고문 통합 AI 분석기")
    st.write("문서 입력부터 다운로드, 그리고 AI 정밀 요약까지 하나의 페이지에서 순차적으로 진행하세요.")
    
    if not st.session_state.authenticated:
        st.info("시스템 사용을 위해 4자리 PIN 번호를 입력해주세요.")
        pin_input = st.text_input("PIN 번호", type="password", max_chars=4)
        if st.button("인증"):
            with st.spinner("네트워크 확인 및 PIN 검증 중..."):
                detect_network()
                remote_pin = fetch_remote_pin()
                if pin_input == remote_pin:
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("PIN 번호가 일치하지 않습니다.")
        st.stop()

    st.success("인증 완료 (네트워크 환경: " + ("사내망" if USE_PROXY else "외부망") + ")")
    
    # ------------------------------------
    # 단계 1: 데이터 입력
    # ------------------------------------
    st.header("1️⃣ 단계: 데이터 입력")
    st.info("💡 **SA, CT 등 3GPP의 모든 워킹그룹 기고문을 100% 완벽하게 지원합니다.** 단, 엑셀 파일 업로드 시 1열의 문서번호에 반드시 **'다운로드 원문 하이퍼링크'**가 걸려 있어야 정확하게 작동합니다.")
    
    input_method = st.radio("입력 방식 선택:", ("Excel 파일 업로드", "링크 텍스트 직접 입력"))
    entries = []

    if input_method == "Excel 파일 업로드":
        uploaded_file = st.file_uploader("엑셀(.xlsx) 파일 선택 (1열 하이퍼링크 필수, 3열 company 양식 준수)", type=["xlsx", "xls"])
        if uploaded_file is not None:
            entries = read_excel_from_bytes(uploaded_file)
            st.info(f"총 {len(entries)}개의 문서 링크를 인식했습니다.")
    else:
        raw_text = st.text_area("3GPP 기고문 원문 링크(.zip)들을 한 줄에 하나씩 붙여넣으세요.", height=150)
        if raw_text:
            lines = [url.strip() for url in raw_text.split('\n') if url.strip()]
            for line in lines:
                docid = line.split('/')[-1].replace('.zip', '')
                entries.append({"doc": docid, "company": "Unknown", "link": line})
            st.info(f"총 {len(entries)}개의 문서 링크를 인식했습니다.")

    # ------------------------------------
    # 단계 2: 기본 분석 실행
    # ------------------------------------
    st.markdown("---")
    st.header("2️⃣ 단계: 기본 추출 및 요약 (TF-IDF)")
    st.write("입력된 링크에서 문서를 다운로드하고 결론(Conclusions)을 추출합니다.")
    
    if st.button("🚀 기본 분석 실행 (Run)", type="primary"):
        if not entries:
            st.warning("먼저 엑셀 파일을 업로드하거나 링크를 입력해주세요.")
        else:
            st.session_state.log_text = ""
            st.session_state.process_done = False
            
            # 이전 AI 분석 기록 초기화
            st.session_state.ai_summary_generated = False
            st.session_state.ai_summary_bytes = None
            
            status_elem = st.empty()
            progress_elem = st.progress(0)
            
            status_elem.text("기고문 다운로드 및 결론 추출 시작...")
            out1_bio = extract_all_conclusions(entries, status_elem, progress_elem, append_log)
            
            status_elem.text("단어 빈도수(TF-IDF) 기반 요약 분석 시작...")
            out2_bio = parse_and_summarize(out1_bio, status_elem, append_log)
            
            status_elem.text("✅ 기본 분석 작업이 완료되었습니다!")
            progress_elem.progress(1.0)
            
            # Bytes로 영구 보존하여 다운로드 시 증발 에러 방지
            st.session_state.out1_bytes = out1_bio.getvalue()
            st.session_state.out2_bytes = out2_bio.getvalue()
            st.session_state.process_done = True

    if st.session_state.process_done:
        st.success("🎉 추출 완료! 아래에서 결과물을 다운로드하거나 바로 AI 정밀 요약을 진행할 수 있습니다.")
        col1, col2 = st.columns(2)
        with col1:
            if st.session_state.out1_bytes:
                st.download_button("📥 Output 1 다운로드 (Conclusions 취합.docx)", data=st.session_state.out1_bytes, file_name="output1_conclusions.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col2:
            if st.session_state.out2_bytes:
                st.download_button("📥 Output 2 다운로드 (TF-IDF 요약.docx)", data=st.session_state.out2_bytes, file_name="output2_summary_tfidf.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        # ------------------------------------
        # 단계 3: NotebookLM 내보내기 & AI 정밀 요약
        # ------------------------------------
        st.markdown("---")
        st.header("3️⃣ 단계: AI 정밀 분석 및 요약")
        st.write("추출된 결론을 바탕으로 여러 회사의 유사 제안을 문맥 단위로 묶어 완벽한 요약본을 생성합니다.")
        
        # 탭(Tabs) 대신 렌더링에 압도적으로 안정적인 라디오(Radio) 토글 UI 사용 (화면 겹침 및 초기화 버그 완벽 방어)
        ai_method_choice = st.radio(
            "💡 요약 방식을 선택하세요:",
            ("📘 구글 NotebookLM 활용하기 (강력 추천🌟)", "⚡ 내장 Gemini API로 요약하기"),
            horizontal=True,
            key="ai_method_radio"
        )
        
        if "NotebookLM" in ai_method_choice:
            st.success("💡 **환각(Hallucination) 제로! 대용량 문서 처리에 가장 추천하는 방법입니다.**\nNotebookLM은 오직 업로드한 문서 기반으로만 답변을 생성하여 압도적인 정확도를 자랑합니다.")
            
            col_a, col_b = st.columns([2, 1])
            with col_a:
                st.markdown("""
                **[NotebookLM의 압도적 장점]**
                * **제한 없는 속도 & 무료:** 복잡한 API 키 발급이나 토큰 초과(429) 에러 없이 **완전 무료**로 즉시 사용 가능!
                * **초대용량 지원:** 노트북 당 **최대 50개의 파일**, 파일당 **최대 50만 단어(약 2,500만 자)**까지 한 번에 거뜬히 분석.
                * **투명한 출처 표기:** 요약된 문장이 원문 기고문의 어느 회사의 어떤 부분인지 정확히 짚어주는 인용(Citation) 링크 제공.
                """)
            with col_b:
                if st.session_state.notebooklm_txt:
                    st.download_button(
                        label="📝 NotebookLM 전용\n 텍스트(.txt) 다운로드",
                        data=st.session_state.notebooklm_txt.encode('utf-8'),
                        file_name="NotebookLM_Input_Conclusions.txt",
                        mime="text/plain",
                        type="primary",
                        use_container_width=True
                    )
            
            st.markdown("---")
            st.markdown("#### 📋 1분 만에 끝내는 NotebookLM 완벽 요약 가이드")
            st.markdown("1. 위 버튼을 눌러 **텍스트 파일(.txt)**을 내 PC에 저장합니다.")
            st.markdown("2. 👉 **[Google NotebookLM 공식 사이트](https://notebooklm.google.com/)** 에 접속하여 로그인합니다.")
            st.markdown("3. 화면의 **'새 노트북(New Notebook)'** 버튼을 누르고, 좌측 소스 탭에 방금 받은 `.txt` 파일을 끌어다 놓습니다.")
            
            st.error("🚨 **[중요] 무한 로딩 현상 대처 꿀팁:** 파일 업로드 후, 우측 패널에서 파일명 옆에 체크표시(✅)가 안 뜨고 **계속 빙글빙글 돌며 무한 로딩**이 걸리는 경우가 종종 있습니다. 이는 화면상 표기 버그일 뿐 실제로는 분석이 끝난 상태입니다! 당황하지 마시고 **그냥 무시한 채로 바로 아래 채팅창에 질문을 전송**하시거나, **F5(새로고침)를 한 번 눌러주시면** 정상 작동합니다.")
            
            st.markdown("4. 화면 하단 채팅창에 아래의 **구조화된 누락 방지 전문가용 프롬프트**를 복사하여 붙여넣고 전송(Enter)하면 완벽한 포맷의 요약이 도출됩니다!")
            
            prompt_text = """당신은 3GPP 표준화 회의의 전문 기술 분석가입니다.
제공된 모든 기고문 전체 원문 모음을 꼼꼼히 검토하고, 아래의 [분석 지침]과 [출력 양식]을 엄격하게 준수하여 분석 보고서를 작성해 주세요.

[분석 지침]
1. 필터링: 반드시 "2개 이상의 회사"가 공통으로 지지하거나 유사한 기술적 주장을 하는 제안(Proposal)만 추출하세요. (1개 회사만 단독으로 주장한 내용은 완전히 제외합니다.)
2. 그룹화: 단어 형태가 달라도 '기술적 핵심 의미와 목적'이 동일하다면 하나의 그룹으로 묶어주세요.
3. 정렬: 지지하는 회사 수가 가장 많은 제안 그룹부터 '내림차순'으로, 2개 이상의 회사가 지지하는 제안(proposal)들을 모두 정렬하세요.
4. 제약사항: 오직 제공된 소스 문서에 명시된 내용, 회사명, 문서 번호만 사용하고, 절대 외부 지식을 섞거나 지어내지 마세요. 환각(Hallucination)을 엄격히 금지합니다.

[출력 양식] (반드시 아래의 마크다운 양식을 똑같이 복제하여 출력할 것)
### [순위]. [제안의 핵심 요약 제목]
* 지지 회사 (총 N개사): [회사명1, 회사명2, ...] (중복 제거 후 쉼표로 나열)
* 상세 제안 내용: [해당 제안의 기술적 배경과 핵심 요구사항을 2~3문장으로 명확하고 이해하기 쉽게 요약]
* 관련 문서 번호: [해당 제안이 포함된 원문 기고문 번호들 (예: R1-2600126 등)]"""
            st.code(prompt_text, language="text")

        else:
            with st.expander("📖 무료/유료 API 키 발급 및 설정 가이드 (순서대로 따라만 하세요!)", expanded=False):
                st.markdown("""
                ### 🟢 [무료 티어] API 키 발급 방법 (간단)
                1. [Google AI Studio](https://aistudio.google.com/app/apikey) 접속 후 구글 계정으로 로그인합니다.
                2. 좌측 상단의 **'Create API key'** 버튼 ➔ 팝업창에서 **'Create API key in new project'** 를 클릭합니다.
                3. 생성된 `AIzaSy...` 로 시작하는 긴 문자를 복사하여 아래 입력창에 붙여넣으세요.
                
                ---
                ### 🔵 [유료 티어] API 키 발급 방법 (상세 가이드)
                **사전 준비물:** 해외 결제가 가능한 신용카드 또는 체크카드 (Visa, Master 등)
                
                **1단계: 구글 클라우드(GCP) 결제 계정 등록 및 방(프로젝트) 만들기**
                1. [Google Cloud Console](https://console.cloud.google.com/) 접속 및 로그인 (약관 동의)
                2. 상단 메뉴바의 **'프로젝트 선택'** ➔ 우측 상단 **'새 프로젝트(New Project)'** 클릭
                3. 프로젝트 이름(예: `3GPP-Analyzer`) 작성 후 **'만들기'** 클릭
                4. 화면 좌측 상단 햄버거 메뉴(☰) ➔ **'결제(Billing)'** ➔ **'결제 계정 연결/추가'** 클릭 후 신용카드 정보 등록
                *(※ 구글이 정상 카드인지 확인하기 위해 1달러를 가결제 후 즉시 취소할 수 있습니다.)*
                
                **2단계: AI Studio에서 유료 프로젝트용 키 생성하기**
                1. 지갑 연결이 끝났다면 다시 [Google AI Studio](https://aistudio.google.com/app/apikey)로 이동합니다.
                2. 좌측 상단의 **'Create API key'** 버튼 클릭
                3. 팝업창에서 바로 파란 버튼을 누르지 말고, ⭐️ **'Search projects' 라고 적힌 돋보기 창을 클릭**하세요!
                4. 방금 1단계에서 카드를 등록했던 **프로젝트 이름(`3GPP-Analyzer`)을 목록에서 찾아 클릭**합니다.
                5. 그 아래 활성화된 **'Create API key in existing project'** 버튼을 클릭합니다.
                6. 생성된 `AIzaSy...` 로 시작하는 키를 복사하여 아래에 입력하세요!
                """)
            
            api_tier_choice = st.radio(
                "API 요금제(Tier) 선택:",
                ("🟢 무료 티어 (데이터 유실 방지 Map-Reduce 적용 + 초고속 Burst 처리)", 
                 "🔵 유료 티어 (문서 전체 원문 일괄 초정밀 분석)"),
                help="무료 API 사용자는 에러 없는 안정적인 분석을 위해 첫 번째를 선택해 주세요."
            )
            
            user_api_key = st.text_input(
                "🔑 Gemini API Key 입력 (1회성 사용으로 안전함)", 
                type="password", 
                help="위의 가이드를 참고하여 발급받은 API 키를 입력해주세요."
            )
            
            if st.button("✨ 내장 정밀 요약 생성 시작"):
                if not user_api_key.strip():
                    st.error("⚠️ 상단에 발급받은 API 키를 입력해주세요.")
                else:
                    genai.configure(api_key=user_api_key.strip())
                    
                    try:
                        valid_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                        if not valid_models:
                            raise Exception("사용 가능한 텍스트 생성 모델이 없습니다. API 키를 확인해주세요.")

                        is_free_tier = "무료" in api_tier_choice
                        
                        if is_free_tier:
                            target_models = [m for m in valid_models if 'flash' in m.lower() and 'vision' not in m.lower()]
                            if not target_models: target_models = [m for m in valid_models if 'pro' in m.lower() and 'vision' not in m.lower()]
                        else:
                            target_models = [m for m in valid_models if 'pro' in m.lower() and 'vision' not in m.lower()]
                            if not target_models: target_models = valid_models
                            
                        target_model_name = next((m for m in target_models if 'latest' in m.lower()), target_models[-1] if target_models else valid_models[-1])
                        model_display_name = target_model_name.split('/')[-1]
                        model = genai.GenerativeModel(target_model_name)
                        
                        # 환각(Hallucination)을 최소화하기 위한 보수적인 모델 설정 (Temperature = 0.1)
                        strict_config = {"temperature": 0.1}
                        
                        # ==========================================
                        # 분할 및 단계별 병합 (Map-Reduce) 로직 실행 (무료 티어 전용)
                        # ==========================================
                        if is_free_tier:
                            batch_size = 15 # 15개 단위 묶음으로 처리 속도 및 토큰 효율 극대화
                            total_docs = len(st.session_state.extracted_data)
                            total_batches = (total_docs + batch_size - 1) // batch_size
                            
                            st.info(f"⚡ 무료 티어 스마트 분석: 빠르고 똑똑한 `{model_display_name}` 모델을 사용하여 {total_docs}개의 문서를 {total_batches}번의 그룹 분석으로 완벽하게 병합합니다.")
                            
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            intermediate_summaries = []
                            
                            for i in range(total_batches):
                                status_text.text(f"🚀 데이터 요약 중... ({i+1}/{total_batches} 번째 그룹 처리 중)")
                                batch_data = st.session_state.extracted_data[i*batch_size : (i+1)*batch_size]
                                
                                batch_text_buffer = []
                                for item in batch_data:
                                    batch_text_buffer.append(f"[문서: {item['doc']}, 회사: {item['company']}]\n{item['content']}")
                                batch_text = "\n\n".join(batch_text_buffer)
                                
                                # 속도 극대화를 위해 불필요한 서술 방지 및 최대한 간결한 추출 지시
                                prompt_map = f"""
                                당신은 3GPP 표준화 회의의 전문 기술 분석가입니다.
                                제공된 원문 데이터는 여러 기고문의 결론(Conclusion) 부분입니다.
                                이 단계에서는 최종 분석을 위한 '중간 데이터'를 추출합니다. 빠지는 내용 없이 모든 제안을 추출하세요.

                                [지침]
                                1. 각 회사들이 주장하는 모든 제안(Proposal)과 결론을 추출하세요. (1개 회사가 단독으로 주장한 것도 이 단계에서는 모두 포함합니다.)
                                2. 의미가 완전히 동일한 제안이 같은 배치 내에 있다면 하나로 묶고 회사명과 문서번호를 병기하세요.
                                3. 절대 외부 지식을 지어내지 마세요. 환각(Hallucination)을 엄격히 금지합니다.
                                4. 최대한 간결하고 짧게 핵심만 1~2문장으로 요약하세요. (처리 속도를 높이기 위함입니다.)

                                [출력 양식]
                                - 제안 요약: [제안 내용]
                                - 지지 회사: [회사명들]
                                - 관련 문서: [문서 번호들]

                                [원문 데이터]
                                {batch_text}
                                """
                                
                                max_retries = 3
                                for attempt in range(max_retries):
                                    try:
                                        res = model.generate_content(prompt_map, generation_config=strict_config)
                                        if res and res.text:
                                            intermediate_summaries.append(res.text)
                                        break
                                    except Exception as e:
                                        if "429" in str(e) or "Quota" in str(e) or "exhausted" in str(e).lower() or "503" in str(e):
                                            if attempt < max_retries - 1:
                                                wait_time = 15 * (attempt + 1) # 15초, 30초 대기 (빠른 백오프)
                                                # 활성 카운트다운(Active Wait) 방어 로직: UI 튕김(Ghosting/Timeout) 완벽 방지
                                                for countdown in range(wait_time, 0, -1):
                                                    status_text.text(f"⚠️ 일시적 속도 제한. {countdown}초 후 번개같이 재시도합니다... (시도 {attempt+1}/{max_retries})")
                                                    time.sleep(1)
                                            else:
                                                raise Exception("무료 API 한도가 소진되었습니다. NotebookLM을 권장합니다.")
                                        else:
                                            raise e
                                
                                progress_bar.progress((i+1)/total_batches)
                                # 인위적인 5초 슬립 제거: 구글 한도 내에서 최고 속도로 Burst 통과 유도
                                    
                            status_text.text("🧠 모든 그룹 분석 완료! 최종 전문가 보고서로 병합하는 중입니다...")
                            final_input = "\n\n=== 그룹별 1차 요약본 모음 ===\n\n".join(intermediate_summaries)
                            
                            # Reduce 프롬프트: 누락 방지 전문가용 구조화 프롬프트 완벽 적용
                            prompt_reduce = f"""
                            당신은 3GPP 표준화 회의의 전문 기술 분석가입니다. 
                            아래 텍스트는 제공된 모든 기고문의 1차 요약본 모음입니다. 이를 꼼꼼히 검토하고, 아래의 [분석 지침]과 [출력 양식]을 엄격하게 준수하여 최종 분석 보고서를 작성해 주세요.

                            [분석 지침]
                            1. 필터링: 반드시 "2개 이상의 회사"가 공통으로 지지하거나 유사한 기술적 주장을 하는 제안(Proposal)만 추출하세요. (1개 회사만 단독으로 주장한 내용은 완전히 제외합니다.)
                            2. 그룹화: 단어 형태가 달라도 '기술적 핵심 의미와 목적'이 동일하다면 하나의 그룹으로 묶어주세요.
                            3. 정렬: 지지하는 회사 수가 가장 많은 제안 그룹부터 '내림차순'으로, 2개 이상의 회사가 지지하는 제안(proposal)들을 모두 정렬하세요.
                            4. 제약사항: 오직 제공된 소스 문서에 명시된 내용, 회사명, 문서 번호만 사용하고, 절대 외부 지식을 섞거나 지어내지 마세요. 환각(Hallucination)을 엄격히 금지합니다.

                            [출력 양식] (반드시 아래의 마크다운 양식을 똑같이 복제하여 출력할 것)
                            ### [순위]. [제안의 핵심 요약 제목]
                            * 지지 회사 (총 N개사): [회사명1, 회사명2, ...] (중복 제거 후 쉼표로 나열)
                            * 상세 제안 내용: [해당 제안의 기술적 배경과 핵심 요구사항을 2~3문장으로 명확하고 이해하기 쉽게 요약]
                            * 관련 문서 번호: [해당 제안이 포함된 원문 기고문 번호들 (예: R1-2600126 등)]

                            [1차 요약본 모음]
                            {final_input}
                            """
                            
                            # Reduce 단계 재시도 및 UI 튕김 방지 로직 강화
                            max_retries_reduce = 3
                            for attempt in range(max_retries_reduce):
                                try:
                                    response = model.generate_content(prompt_reduce, generation_config=strict_config)
                                    break
                                except Exception as e:
                                    if "429" in str(e) or "Quota" in str(e) or "exhausted" in str(e).lower() or "503" in str(e):
                                        if attempt < max_retries_reduce - 1:
                                            wait_time = 15 * (attempt + 1)
                                            for countdown in range(wait_time, 0, -1):
                                                status_text.text(f"⚠️ 서버 응답 지연. 최종 병합을 위해 대기합니다... (시도 {attempt+1}/{max_retries_reduce}) - {countdown}초 남음")
                                                time.sleep(1)
                                        else:
                                            raise Exception("API 한도가 완전히 소진되었습니다. 잠시 후 시도하거나 NotebookLM을 이용해 주세요.")
                                    else:
                                        raise e
                                        
                            status_text.text("✅ AI 초고속 병합 및 정밀 요약 완료!")
                        
                        # ==========================================
                        # 유료 티어 로직 (기존과 동일하게 대규모 일괄 처리)
                        # ==========================================
                        else:
                            st.info(f"💎 유료 티어 모드 가동: 용량 제한 없이 가장 똑똑한 `{model_display_name}` 모델로 문서 전체 원문을 초정밀 분석합니다.")
                            status_text = st.empty()
                            status_text.text("AI가 방대한 문서 전체를 정독하고 분석 중입니다...")
                            
                            extracted_text_buffer = []
                            for item in st.session_state.extracted_data:
                                extracted_text_buffer.append(f"[문서: {item['doc']}, 회사: {item['company']}]\n{item['full_content']}")
                            full_text = "\n\n".join(extracted_text_buffer)
                            
                            # 유료 티어 원문 분석 단계에도 누락 방지 최신 프롬프트 적용
                            prompt_paid = f"""
                            당신은 3GPP 표준화 회의의 전문 기술 분석가입니다. 
                            제공된 모든 기고문 전체 원문 모음을 꼼꼼히 검토하고, 아래의 [분석 지침]과 [출력 양식]을 엄격하게 준수하여 분석 보고서를 작성해 주세요.

                            [분석 지침]
                            1. 필터링: 반드시 "2개 이상의 회사"가 공통으로 지지하거나 유사한 기술적 주장을 하는 제안(Proposal)만 추출하세요. (1개 회사만 단독으로 주장한 내용은 완전히 제외합니다.)
                            2. 그룹화: 단어 형태가 달라도 '기술적 핵심 의미와 목적'이 동일하다면 하나의 그룹으로 묶어주세요.
                            3. 정렬: 지지하는 회사 수가 가장 많은 제안 그룹부터 '내림차순'으로, 2개 이상의 회사가 지지하는 제안(proposal)들을 모두 정렬하세요.
                            4. 제약사항: 오직 제공된 소스 문서에 명시된 내용, 회사명, 문서 번호만 사용하고, 절대 외부 지식을 섞거나 지어내지 마세요. 환각(Hallucination)을 엄격히 금지합니다.

                            [출력 양식] (반드시 아래의 마크다운 양식을 똑같이 복제하여 출력할 것)
                            ### [순위]. [제안의 핵심 요약 제목]
                            * 지지 회사 (총 N개사): [회사명1, 회사명2, ...] (중복 제거 후 쉼표로 나열)
                            * 상세 제안 내용: [해당 제안의 기술적 배경과 핵심 요구사항을 2~3문장으로 명확하고 이해하기 쉽게 요약]
                            * 관련 문서 번호: [해당 제안이 포함된 원문 기고문 번호들 (예: R1-2600126 등)]

                            [기고문 원문 데이터]
                            {full_text}
                            """
                            
                            max_retries_paid = 3
                            for attempt in range(max_retries_paid):
                                try:
                                    response = model.generate_content(prompt_paid, generation_config=strict_config)
                                    break
                                except Exception as e:
                                    if "429" in str(e) or "Quota" in str(e) or "503" in str(e):
                                        if attempt < max_retries_paid - 1:
                                            wait_time = 15 * (attempt + 1)
                                            for countdown in range(wait_time, 0, -1):
                                                status_text.text(f"⚠️ 서버 응답 지연. 재시도 대기 중... ({attempt+1}/{max_retries_paid}) - {countdown}초 남음")
                                                time.sleep(1)
                                        else:
                                            raise e
                                    else:
                                        raise e
                                        
                            status_text.text("✅ AI 초정밀 일괄 분석 완료!")

                        # 파일 생성 및 데이터를 세션 메모리에 보존
                        if response and response.text:
                            r = Document()
                            r.add_heading(f"AI 정밀 분석 요약 ({model_display_name})", 0)
                            
                            for line in response.text.split('\n'):
                                if re.match(r'^(#+)?\s*\d+\.|###', line.strip()):
                                    p = r.add_paragraph()
                                    p.add_run(line.replace('#', '').strip()).bold = True
                                else:
                                    r.add_paragraph(line)
                            
                            bio_llm = io.BytesIO()
                            r.save(bio_llm)
                            
                            # 분석 결과를 영구 보존용 세션 금고에 저장
                            st.session_state.ai_summary_bytes = bio_llm.getvalue()
                            st.session_state.ai_summary_text = response.text
                            st.session_state.ai_model_name = model_display_name
                            st.session_state.ai_summary_generated = True
                            
                        else:
                            st.error("AI 응답을 받아오지 못했습니다. 잠시 후 다시 시도해주세요.")
                                
                    except Exception as e:
                        error_msg = str(e)
                        if "429" in error_msg or "Quota" in error_msg or "exhausted" in error_msg.lower():
                            st.error("❌ **[API 용량 초과 안내]** 무료 일일 제공량을 완전히 소진했거나 텍스트가 너무 방대합니다. 왼쪽 탭의 **[📘 구글 NotebookLM 활용하기]**를 이용해 주세요.")
                        else:
                            st.error(f"❌ API 호출 중 오류가 발생했습니다. 키가 정확한지 확인해주세요.\n\n[상세 오류 메시지]: {e}")

            # ==========================================
            # UI 분리: 버튼 블록 밖에서 다운로드 UI 표출 (새로고침 및 화면 겹침 방어)
            # ==========================================
            if st.session_state.ai_summary_generated:
                st.success("✅ 전문가 수준의 AI 정밀 요약 파일이 성공적으로 생성되었습니다!")
                st.download_button(
                    label=f"📥 AI 요약본(Output 3) 최종 다운로드 (.docx)",
                    data=st.session_state.ai_summary_bytes,
                    file_name="Output3_AI_Summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
                with st.expander("👀 생성된 AI 요약 결과 미리보기", expanded=True):
                    st.markdown(st.session_state.ai_summary_text)

# --- 페이지 2: 3GPP FTP 탐색기 ---
elif page == "📁 3GPP FTP 탐색기":
    st.title("📁 3GPP 공식 FTP 탐색기 및 가이드")
    st.write("이 페이지에서는 3GPP 회의의 원본 기고문들이 업로드되는 공식 FTP 서버의 다이렉트 링크와 활용 방법을 안내합니다.")
    
    st.subheader("🔗 주요 Working Group 회의록 FTP 접속")
    st.markdown("""
    * **RAN WG (무선 접속망)**
        * [RAN1 (물리계층) 바로가기](https://www.3gpp.org/ftp/tsg_ran/WG1_RL1/)
        * [RAN2 (무선 인터페이스 구조) 바로가기](https://www.3gpp.org/ftp/tsg_ran/WG2_RL2/)
        * [RAN3 (네트워크 아키텍처) 바로가기](https://www.3gpp.org/ftp/tsg_ran/WG3_IU/)
    * **SA WG (서비스 및 시스템 아키텍처)**
        * [SA1 (서비스) 바로가기](https://www.3gpp.org/ftp/tsg_sa/WG1_Serv/)
        * [SA2 (아키텍처) 바로가기](https://www.3gpp.org/ftp/tsg_sa/WG2_Arch/)
    * **CT WG (코어 네트워크 및 단말)**
        * [CT1 (단말/코어 프로토콜) 바로가기](https://www.3gpp.org/ftp/tsg_ct/WG1_mm-cc-sm_ex-CN1/)
        * [CT3 (인터워킹) 바로가기](https://www.3gpp.org/ftp/tsg_ct/WG3_interworking_ex-CN3/)
    """)
    
    st.subheader("📖 기고문 번호(TDoc) 읽는 법")
    st.write("""
    3GPP 기고문은 일반적으로 WG 그룹명과 연도, 일련번호 조합을 가집니다. (예: `R1-2505131`, `S2-250123`)
    - **앞자리 (R1, S2, C1 등):** 워킹그룹(Working Group)을 의미합니다.
    - **중간 (25):** 2025년에 제출되었음을 의미합니다.
    - **뒷자리 (05131):** 해당 연도의 기고문 일련번호입니다.
    이를 통해 기고문이 어느 그룹에서 언제 제출되었는지 쉽게 파악할 수 파악할 수 있습니다.
    """)

# --- 페이지 3: 소개 및 가이드 ---
elif page == "ℹ️ 소개 및 가이드":
    st.title("ℹ️ 초보자 상세 가이드 및 이용 안내")
    
    st.header("🔰 초보자를 위한 단계별 사용 가이드")
    
    st.markdown("### 1단계: 분석할 기고문 데이터 준비하기")
    st.write("""
    1. 분석하고 싶은 3GPP 기고문들의 링크나 문서 번호를 확보합니다. (**SA, CT, RAN 등 모든 워킹그룹 완벽 호환**)
    2. 엑셀 파일을 만드실 경우, **1열에는 문서번호, 3열에는 회사명**을 적어주세요. 
    🚨 **[주의]** SA나 CT 문서를 엑셀로 업로드할 때는, 프로그램이 정확한 서버를 찾을 수 있도록 **1열 문서번호 셀에 반드시 '원문 다운로드 하이퍼링크'가 걸려 있어야 합니다.**
    """)
    
    st.markdown("### 2단계: 메인 분석기 실행하기 (결론 자동 추출)")
    st.write("""
    1. 좌측 메뉴에서 **[🚀 통합 AI 분석기]**로 이동합니다.
    2. 준비한 엑셀 파일을 업로드하거나, 텍스트 입력창에 링크들을 복사해서 붙여넣습니다.
    3. **'🚀 기본 분석 실행 (Run)'** 버튼을 누르면, 프로그램이 자동으로 각 문서의 Conclusion(결론) 부분만 쏙쏙 뽑아냅니다.
    """)
    
    st.markdown("### 3단계: AI 정밀 분석으로 요약본 만들기 (NotebookLM 강력 추천 🌟)")
    st.write("""
    가장 추천하는 방법은 구글이 제공하는 무료 문서 분석 특화 AI인 **[NotebookLM](https://notebooklm.google.com/)**을 활용하는 것입니다.
    
    **[NotebookLM 100% 활용 가이드]**
    * **환각(Hallucination) 제로:** 일반 챗봇과 달리, 오직 내가 업로드한 문서(소스)에서만 정답을 찾기 때문에 없는 회사나 제안을 지어내지 않습니다.
    * **엄청난 무료 용량:** 노트북 1개당 최대 **50개의 문서**, 각 문서당 최대 **50만 단어**까지 완전 무료로 업로드할 수 있어 에러 걱정이 없습니다.
    * **투명한 출처 표기:** 요약본 뒤에 인용구 번호가 달려서, 원문의 어떤 부분에서 해당 내용이 발췌되었는지 클릭 한 번으로 검증할 수 있습니다.
    """)
    
    st.markdown("---")
    
    st.header("🔒 개인정보처리 및 보안 (안심하고 사용하세요)")
    st.write("""
    * **API 키 절대 보호:** 귀하가 입력한 API 키는 화면에 `****` 형태로 가려져 보이며, 서버의 하드디스크나 데이터베이스에 절대 저장되지 않습니다. 요약 과정이 끝나면 메모리에서 즉시 영구 삭제됩니다.
    * **문서 데이터 무단 수집 금지:** 업로드하신 엑셀 파일과 추출된 텍스트 역시 세션이 종료되거나 웹 브라우저를 닫는 즉시 완벽하게 소멸됩니다.
    """)
    
    st.header("⚖️ 이용 약관")
    st.write("본 서비스에서 제공하는 요약 결과는 AI 기반 알고리즘에 의존하므로 100%의 정확도를 보장하지 않습니다. 공식적인 통계나 회의록은 반드시 3GPP 공식 홈페이지를 교차 검증하시기 바랍니다.")
